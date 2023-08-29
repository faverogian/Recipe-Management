'''
Recipe Manager
Docx Output Class
Gian Favero
2022
'''

# Python Imports
from docx import Document
from docx.shared import Pt
from datetime import date

def createDocument(shoppingList, menu):
    document = Document()

    today = date.today()

    shoppingTitle = document.add_heading(f"Shopping List: Week of {today.strftime('%A')}, {today.strftime('%B')} {today.strftime('%d')}, {today.strftime('%Y')}", 0)
    shoppingTitle.alignment = 1

    keys = ["Meats", "Poultry", "Dairy", "Vegetables", "Bakery", "Pastas", "Spices", "Snacks", "Other"]

    for key in keys:
        keyText = document.add_paragraph().add_run(f"{key}")
        keyText.font.size = Pt(14)
        para = document.add_paragraph()
        try:
            para.add_run(', '.join(shoppingList[key])) 
        except:
            pass

    document.add_page_break()

    menuTitle = document.add_heading(f"Menu: Week of {today.strftime('%A')}, {today.strftime('%B')} {today.strftime('%d')}, {today.strftime('%Y')}", 0)
    menuTitle.alignment = 1

    days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"]

    for day in days:
        dayText = document.add_paragraph().add_run(f"{day}")
        dayText.font.size = Pt(14)
        para = document.add_paragraph()
        try:
            para.add_run(', '.join(menu[day]))
        except:
            pass

    document.save(r"Outputs/This_Week_In_Food.docx")
